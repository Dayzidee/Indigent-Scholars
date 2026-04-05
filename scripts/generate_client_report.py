#!/usr/bin/env python3
"""
Indigent-Scholars: Client Progress Report Generator
Generates a professional Word document summarizing development progress
with embedded Stitch UI design screenshots.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# ──────────────────────────────────────────────
# Configuration
# ──────────────────────────────────────────────
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
STITCH_DIR = os.path.join(PROJECT_ROOT, "src", "stitch-exports", "screens")
OUTPUT_PATH = os.path.join(PROJECT_ROOT, "Indigent-Scholars_Client_Report.docx")
DATE_STR = datetime.now().strftime("%B %d, %Y")

# Categorised screens with human-readable names and descriptions
SCREEN_CATEGORIES = {
    "Public-Facing Pages (Completed & Coded)": {
        "description": "These pages are fully designed, coded in Next.js with Tailwind CSS, and feature entrance animations powered by Framer Motion.",
        "screens": [
            {
                "folder": "Home_-_Minimalist_Power__Final_Refinement_",
                "title": "Homepage — Minimalist Power",
                "desc": "The main landing page featuring a 3D interactive sphere background, hero section with call-to-action, impact statistics, and testimonials. Fully responsive with scroll-driven animations."
            },
            {
                "folder": "About_Us_-_The_Vision___Mission",
                "title": "About Us — Vision & Mission",
                "desc": "Showcases the organisation's story, mission, and the team behind EdAfrica. Features staggered fade-in animations and editorial typography."
            },
            {
                "folder": "Verified_Students_Gallery",
                "title": "Verified Students Gallery",
                "desc": "A public gallery of verified students seeking sponsorship. Designed for sponsors to browse and connect with scholars."
            },
            {
                "folder": "For_Sponsors_-_Partnering_for_Impact",
                "title": "For Sponsors — Partnering for Impact",
                "desc": "A dedicated page for potential sponsors explaining the process, impact metrics, and how to get involved."
            },
        ]
    },
    "Authentication & Onboarding (Completed & Coded)": {
        "description": "User authentication flows are designed and coded. Backend integration with Supabase Auth is in progress.",
        "screens": [
            {
                "folder": "Log_In_-_Indigent-Sc",
                "title": "Login Page",
                "desc": "Clean, minimal login interface with email/password fields and social sign-in options. Features glassmorphism card styling."
            },
            {
                "folder": "Sponsor_Registration_Form",
                "title": "Sponsor Registration",
                "desc": "Multi-step sponsor onboarding form capturing organisation details, contact information, and verification documents."
            },
            {
                "folder": "Reset_Your_Password",
                "title": "Password Reset",
                "desc": "Password recovery flow with email verification step."
            },
        ]
    },
    "Student Dashboard (Designed — Development Next Phase)": {
        "description": "Complete UI designs for the student portal. These will be coded in the next development sprint and connected to the Supabase database.",
        "screens": [
            {
                "folder": "Student_Dashboard_-_Profile_Incomplete",
                "title": "Student Dashboard — Profile (Incomplete State)",
                "desc": "The initial state a student sees after registration, prompting them to complete their profile and upload documents."
            },
            {
                "folder": "Student_Dashboard_-_My_Application",
                "title": "Student Dashboard — My Application",
                "desc": "Application tracking interface showing submission status (Pending, Verified, Rejected) with document upload capabilities."
            },
            {
                "folder": "Student_Dashboard_-_Verified_State",
                "title": "Student Dashboard — Verified State",
                "desc": "The verified student view showing their public profile status, funding received, and sponsor matches."
            },
            {
                "folder": "Student_Dashboard_-_Sponsor_Matches__Locked_",
                "title": "Student Dashboard — Sponsor Matches (Locked)",
                "desc": "Shows pending sponsor matches that unlock after admin verification is complete."
            },
            {
                "folder": "Student_Dashboard_-_Settings",
                "title": "Student Dashboard — Settings",
                "desc": "Account settings page for students to manage their profile, notifications, and privacy preferences."
            },
        ]
    },
    "Sponsor Dashboard (Designed — Development Next Phase)": {
        "description": "Complete UI designs for the sponsor portal. Scheduled for development after the student dashboard.",
        "screens": [
            {
                "folder": "Sponsor_Dashboard_-_Impact_Overview",
                "title": "Sponsor Dashboard — Impact Overview",
                "desc": "The sponsor's main dashboard showing total funded amount, number of scholars supported, and impact metrics."
            },
            {
                "folder": "Sponsor_Dashboard_-_Discover_Students",
                "title": "Sponsor Dashboard — Discover Students",
                "desc": "A curated directory of verified students available for sponsorship, with filtering and search capabilities."
            },
            {
                "folder": "Sponsor_Dashboard_-_Funding_Ledger",
                "title": "Sponsor Dashboard — Funding Ledger",
                "desc": "Detailed transaction history showing all pledges, fulfilled commitments, and payment records."
            },
            {
                "folder": "Sponsor_Dashboard_-_Pending_Verification_Teaser",
                "title": "Sponsor Dashboard — Pending Verification",
                "desc": "Pre-verification state showing the sponsor that their account is under review before they can begin funding."
            },
            {
                "folder": "Sponsor_Dashboard_-_Org_Settings",
                "title": "Sponsor Dashboard — Organisation Settings",
                "desc": "Organisation profile management for sponsors including logo upload, company details, and team access."
            },
        ]
    },
    "Admin Portal (Designed — Development Next Phase)": {
        "description": "Complete UI designs for the administrative verification and management portal.",
        "screens": [
            {
                "folder": "Admin_Dashboard_-_Global_Overview",
                "title": "Admin Dashboard — Global Overview",
                "desc": "Platform-wide analytics showing total users, pending verifications, active funding matches, and key performance indicators."
            },
            {
                "folder": "Admin_-_Verification_Center__Student_List_",
                "title": "Admin — Verification Center (Student List)",
                "desc": "Queue of student applications awaiting verification, with quick-action buttons to approve or reject."
            },
            {
                "folder": "Admin_-_Verification_Center__Sponsor_List_",
                "title": "Admin — Verification Center (Sponsor List)",
                "desc": "Queue of sponsor registrations awaiting verification with organisation details at a glance."
            },
            {
                "folder": "Admin_-_Verification_Center__Split_View_",
                "title": "Admin — Verification Center (Split View)",
                "desc": "Side-by-side document review interface allowing admins to inspect uploaded documents and approve/reject in one view."
            },
            {
                "folder": "Admin_-_User_CRM_with_Profile_Drawer",
                "title": "Admin — User CRM with Profile Drawer",
                "desc": "Full user management CRM with a slide-out drawer for detailed profile views and administrative actions."
            },
            {
                "folder": "Admin_-_Financial_Command_Ledger",
                "title": "Admin — Financial Command Ledger",
                "desc": "Platform-wide financial overview of all funding matches, pledges, and disbursements."
            },
            {
                "folder": "Admin_-_Platform_Management_Settings",
                "title": "Admin — Platform Management Settings",
                "desc": "System-wide configuration for roles, permissions, email templates, and platform policies."
            },
        ]
    },
    "Additional Screens": {
        "description": "Supporting screens for the user experience.",
        "screens": [
            {
                "folder": "Student_Profile_Detail",
                "title": "Student Profile Detail",
                "desc": "Detailed public profile view of a verified student, visible to sponsors considering funding."
            },
        ]
    }
}


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def add_styled_heading(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0x3D, 0x9B)  # Trust Blue
    return heading


def build_document():
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1C, 0x1B, 0x1B)

    # ════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════
    # Add some spacing before the title
    for _ in range(4):
        doc.add_paragraph("")

    # Logo if available
    logo_path = os.path.join(STITCH_DIR, "cropped-lite-logo-300x62.png", "screenshot.png")
    if os.path.exists(logo_path):
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.add_run().add_picture(logo_path, width=Inches(3))

    doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INDIGENT-SCHOLARS")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0x3D, 0x9B)
    run.font.name = 'Calibri'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Platform Development Progress Report")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x5F, 0x5E, 0x5E)

    doc.add_paragraph("")

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(DATE_STR)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x73, 0x76, 0x85)

    doc.add_paragraph("")

    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = confidential.add_run("CONFIDENTIAL — For Client Review Only")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xBA, 0x1A, 0x1A)

    doc.add_page_break()

    # ════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ════════════════════════════════════════════
    add_styled_heading(doc, "1. Executive Summary", level=1)

    doc.add_paragraph(
        "This report provides a comprehensive update on the development progress of the "
        "Indigent-Scholars platform — a philanthropic EdTech web application designed to connect "
        "Nigerian scholars in financial need with sponsors and donors worldwide."
    )

    doc.add_paragraph(
        "The platform is being built using a modern, production-grade technology stack and follows "
        "a premium \"Scholarly Architect\" design system to ensure the interface conveys institutional "
        "trust and warmth."
    )

    doc.add_paragraph("")

    # ── Status Table ──
    add_styled_heading(doc, "Current Status at a Glance", level=2)

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Shading Accent 1'

    headers = ["Area", "Status"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

    data = [
        ("Public Pages (Home, About, Students, Sponsors)", "✅ Completed & Coded"),
        ("Authentication UI (Login, Register, Reset)", "✅ Completed & Coded"),
        ("3D Interactive Background (Three.js)", "✅ Completed & Coded"),
        ("Dashboard Designs (Student, Sponsor, Admin)", "🎨 Designed — Ready for Development"),
        ("Backend & Database (Supabase)", "🔧 Architecture Defined — Setup In Progress"),
    ]

    for row_idx, (area, status) in enumerate(data, start=1):
        table.rows[row_idx].cells[0].text = area
        table.rows[row_idx].cells[1].text = status

    doc.add_paragraph("")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # TECHNOLOGY STACK
    # ════════════════════════════════════════════
    add_styled_heading(doc, "2. Technology Stack", level=1)

    tech_items = [
        ("Frontend Framework", "Next.js 15 (React 19) with App Router"),
        ("Styling", "Tailwind CSS v4 with custom \"Scholarly Architect\" design tokens"),
        ("Animations", "Framer Motion (UI transitions) + GSAP ScrollTrigger (scroll-driven effects)"),
        ("3D Graphics", "Three.js via @react-three/fiber and @react-three/drei"),
        ("Backend / Database", "Supabase (PostgreSQL with Row Level Security)"),
        ("Authentication", "Supabase Auth (Email/Password + OAuth)"),
        ("File Storage", "Supabase Storage (for student documents)"),
        ("Language", "TypeScript (strict mode)"),
    ]

    tech_table = doc.add_table(rows=len(tech_items) + 1, cols=2)
    tech_table.style = 'Light Shading Accent 1'

    tech_table.rows[0].cells[0].text = "Component"
    tech_table.rows[0].cells[1].text = "Technology"
    for p in tech_table.rows[0].cells[0].paragraphs:
        for run in p.runs:
            run.font.bold = True
    for p in tech_table.rows[0].cells[1].paragraphs:
        for run in p.runs:
            run.font.bold = True

    for i, (comp, tech) in enumerate(tech_items, start=1):
        tech_table.rows[i].cells[0].text = comp
        tech_table.rows[i].cells[1].text = tech

    doc.add_paragraph("")
    doc.add_page_break()

    # ════════════════════════════════════════════
    # DESIGN SCREENS BY CATEGORY
    # ════════════════════════════════════════════
    add_styled_heading(doc, "3. UI/UX Design & Development Progress", level=1)

    doc.add_paragraph(
        "Below is a categorised overview of all designed screens, with screenshots from our "
        "design system (Stitch by Google). Screens marked as \"Completed & Coded\" are already "
        "implemented in the production codebase. Screens marked as \"Designed\" are approved "
        "mockups ready for the next development sprint."
    )
    doc.add_paragraph("")

    section_num = 1
    for category, info in SCREEN_CATEGORIES.items():
        add_styled_heading(doc, f"3.{section_num}  {category}", level=2)

        # Category description
        desc_para = doc.add_paragraph()
        run = desc_para.add_run(info["description"])
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x73, 0x76, 0x85)
        doc.add_paragraph("")

        for screen in info["screens"]:
            # Screen title
            screen_heading = doc.add_paragraph()
            run = screen_heading.add_run(f"▸ {screen['title']}")
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0x52, 0xCC)

            # Screen description
            doc.add_paragraph(screen["desc"])

            # Screenshot image
            img_path = os.path.join(STITCH_DIR, screen["folder"], "screenshot.png")
            if os.path.exists(img_path):
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    img_para.add_run().add_picture(img_path, width=Inches(5.5))
                except Exception as e:
                    run = img_para.add_run(f"[Image could not be embedded: {e}]")
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0xBA, 0x1A, 0x1A)
            else:
                missing = doc.add_paragraph()
                run = missing.add_run(f"[Screenshot not found: {screen['folder']}]")
                run.font.italic = True
                run.font.color.rgb = RGBColor(0xBA, 0x1A, 0x1A)

            doc.add_paragraph("")  # spacing

        section_num += 1
        doc.add_page_break()

    # ════════════════════════════════════════════
    # DATABASE ARCHITECTURE
    # ════════════════════════════════════════════
    add_styled_heading(doc, "4. Database Architecture", level=1)

    doc.add_paragraph(
        "The backend is powered by Supabase (hosted PostgreSQL). The database schema has been "
        "fully designed and is being deployed. It includes the following core tables:"
    )

    db_items = [
        ("Profiles", "Extends Supabase Auth users with role (student/sponsor/admin), full name, and email."),
        ("Student Applications", "Stores student submissions including school info, funding needs, uploaded documents, and verification status (pending/verified/rejected)."),
        ("Sponsors", "Stores sponsor profiles including organisation name, contact details, total funded amount, and verification status."),
        ("Funding Matches", "Tracks sponsor-to-student funding commitments with pledge status (pledged/fulfilled/cancelled)."),
    ]

    for name, desc in db_items:
        para = doc.add_paragraph()
        run = para.add_run(f"• {name}: ")
        run.font.bold = True
        para.add_run(desc)

    doc.add_paragraph("")

    doc.add_paragraph(
        "All tables are protected by Row Level Security (RLS) policies ensuring that students "
        "can only access their own data, sponsors can only view verified students, and admins "
        "have full oversight capabilities."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════
    # SECURITY
    # ════════════════════════════════════════════
    add_styled_heading(doc, "5. Security & Access Control", level=1)

    security_items = [
        "Row Level Security (RLS) on every database table — users can only access data they are authorised to see.",
        "Next.js Edge Middleware enforcing authentication on all dashboard routes — unauthenticated users are redirected to the login page.",
        "Role-Based Access Control (RBAC) — students, sponsors, and admins each see only their designated portal.",
        "Supabase Auth handles password hashing, session management, and token refresh securely.",
        "Document uploads are stored in a dedicated Supabase Storage bucket with access policies.",
    ]

    for item in security_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ════════════════════════════════════════════
    # NEXT STEPS
    # ════════════════════════════════════════════
    add_styled_heading(doc, "6. Next Steps & Roadmap", level=1)

    phases = [
        {
            "name": "Phase 1 — Database Deployment (Current)",
            "items": [
                "Deploy the PostgreSQL schema to the live Supabase project.",
                "Configure storage bucket for student document uploads.",
                "Set up authentication triggers for automatic profile creation on sign-up.",
            ]
        },
        {
            "name": "Phase 2 — Dashboard Development",
            "items": [
                "Build the Student Dashboard (profile management, application tracking, document uploads).",
                "Build the Sponsor Dashboard (impact overview, scholar directory, funding commitments).",
                "Build the Admin Portal (verification queue, user CRM, financial ledger).",
            ]
        },
        {
            "name": "Phase 3 — Integration & Testing",
            "items": [
                "Connect all dashboard components to the live Supabase database.",
                "Implement real-time data updates for status changes.",
                "End-to-end testing of all user flows (student, sponsor, admin).",
            ]
        },
        {
            "name": "Phase 4 — Polish & Deployment",
            "items": [
                "Performance optimisation (image compression, lazy loading).",
                "SEO and accessibility audit.",
                "Production deployment and domain configuration.",
            ]
        },
    ]

    for phase in phases:
        phase_para = doc.add_paragraph()
        run = phase_para.add_run(phase["name"])
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0x52, 0xCC)

        for item in phase["items"]:
            doc.add_paragraph(item, style='List Bullet')
        doc.add_paragraph("")

    doc.add_page_break()

    # ════════════════════════════════════════════
    # CLOSING
    # ════════════════════════════════════════════
    add_styled_heading(doc, "7. Summary", level=1)

    doc.add_paragraph(
        "The Indigent-Scholars platform is progressing on schedule. All public-facing pages are "
        "fully developed and feature premium animations and a bespoke design system. The complete "
        "dashboard suite (Student, Sponsor, and Admin portals) has been designed and approved, "
        "with development commencing immediately. The database architecture is finalised and "
        "deployment is underway."
    )

    doc.add_paragraph("")

    doc.add_paragraph(
        "We are committed to delivering a world-class platform that empowers Nigerian scholars "
        "and connects them with sponsors who want to make a lasting impact in education."
    )

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Footer note
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— End of Report —")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x73, 0x76, 0x85)

    doc.add_paragraph("")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run("For questions or feedback, please contact the development team.")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x73, 0x76, 0x85)

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"\n{'='*60}")
    print(f"  ✅ Client report generated successfully!")
    print(f"  📄 File: {OUTPUT_PATH}")
    print(f"  📅 Date: {DATE_STR}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    build_document()
